import os
from flask import Flask, request, render_template, jsonify, send_from_directory, redirect, url_for
import openai
import numpy as np
import pandas as pd
import json
import re
from json import JSONDecodeError
from typing import Dict, List, Union
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
import warnings
from datetime import datetime
from werkzeug.utils import secure_filename
import threading
import logging
import mammoth

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")

# 初始化 Flask 应用
app = Flask(__name__)

# 配置路径
UPLOAD_FOLDER = 'uploads'
REPORT_FOLDER = 'reports'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['REPORT_FOLDER'] = REPORT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB 文件大小限制
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['REPORT_FOLDER'], exist_ok=True)

# OpenAI 配置
openai.api_key = "sk-2aJtjc7nVJZMNMRk2kHpQrQWUphatN7YwurCqew18"
openai.api_base = "https://free.yunwu.ai/v1"

# 全局变量存储进度
progress = {"status": "idle", "percentage": 0, "step": "等待上传"}
conversation_history = []

# JSON 输出格式模板
RESPONSE_FORMATS = {
    "sensitivity_standards": {"format": """[{"field": "字段名", "sensitivity": "高/中/低", "reason": "理由说明"}]"""},
    "ahp_matrix": {
        "format": """{"matrix": [[1,2,3],[0.5,1,2],[0.33,0.5,1]], "indicators": [{"name": "指标名称", "explanation": "详细解释"}]}"""},
    "scoring": {"format": """[{"field": "字段名", "scores": {"指标1": 分数, "指标2": 分数}}]"""}
}


# API 调用函数
def get_response(prompt: str, step_name: str) -> str:
    global progress
    full_prompt = f"{prompt}\n请严格使用以下JSON格式输出：{RESPONSE_FORMATS[step_name]['format']}"
    conversation_history.append({"role": "user", "content": full_prompt})
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=conversation_history
        )
        reply = response.choices[0].message["content"]
        conversation_history.append({"role": "assistant", "content": reply})
        logger.info(f"API 调用成功: {step_name}")
        progress["step"] = f"完成 {step_name} 调用"
        return reply
    except Exception as e:
        logger.error(f"API错误: {e}")
        progress["step"] = f"API 错误: {str(e)}"
        return None


# JSON 解析函数
def parse_response(response: str, expected_type: str) -> Union[dict, list]:
    def clean_json_string(s):
        s = re.sub(r'[\x00-\x1F]+', '', s)
        s = re.sub(r'//.*', '', s)
        s = re.sub(r'/\*.*?\*/', '', s, flags=re.DOTALL)
        return s.strip()

    def fix_common_issues(s):
        s = re.sub(r',(\s*[}\]])', r'\1', s)
        s = re.sub(r"(\w+):", r'"\1":', s)
        s = re.sub(r"'(.*?)'", r'"\1"', s)
        return s

    try:
        json_match = re.search(r'(\[.*\]|\{.*\})', response, re.DOTALL) or \
                     re.search(r'```(?:json)?\s*(\{.*\}|$$ .* $$)\s*```', response, re.DOTALL)
        if json_match:
            raw_content = json_match.group(1)
            cleaned = clean_json_string(raw_content)
            repaired = fix_common_issues(cleaned)
            return json.loads(repaired)
        raise ValueError("未检测到JSON结构")
    except JSONDecodeError as e:
        logger.error(f"JSON解析失败: {e}")
        if e.msg == 'Extra data':
            repaired = repaired[:e.pos]
            return json.loads(repaired)
        raise ValueError("无法解析JSON")


# 评分数据处理
def process_scoring_data(data: List[dict], indicators: List[str]) -> pd.DataFrame:
    from difflib import get_close_matches
    validated = []
    for item in data:
        scores = item.get('scores', {})
        item_scores = {}
        for ind in indicators:
            matches = get_close_matches(ind, scores.keys(), n=1, cutoff=0.6)
            score = scores[matches[0]] if matches else 0
            if isinstance(score, (int, float)) and 1 <= score <= 10:
                item_scores[ind] = score
            else:
                item_scores[ind] = 0
        validated.append({"field": item.get("field", "未知字段"), **item_scores})
    if not validated:
        raise ValueError("未解析到有效评分数据")
    df = pd.DataFrame(validated).set_index("field")
    return df[indicators]


# AHP 计算
def calculate_weights(matrix: np.ndarray) -> np.ndarray:
    row_products = np.prod(matrix, axis=1) ** (1 / matrix.shape[1])
    return row_products / np.sum(row_products)


def consistency_check(matrix: np.ndarray) -> tuple:
    n = matrix.shape[0]
    ri_table = [0, 0, 0.58, 0.90, 1.12, 1.24, 1.32, 1.41, 1.45, 1.49]
    if n > len(ri_table) - 1:
        raise ValueError("矩阵维度超过RI表范围")
    eigenvalues = np.linalg.eigvals(matrix)
    max_eigen = np.max(eigenvalues.real)
    ci = (max_eigen - n) / (n - 1)
    cr = ci / ri_table[n]
    return cr < 0.1, cr


# 分析和报告生成函数
def analyze_and_generate_report(file_path, excel_report_path, word_report_path):
    global conversation_history, progress
    conversation_history = []
    progress["status"] = "running"
    progress["percentage"] = 10
    progress["step"] = "读取文件"
    logger.info(f"开始分析文件: {file_path}")

    try:
        # 读取 CSV
        df = pd.read_csv(file_path)
        fields = df.columns.tolist()
        sample_data = df.head(5).to_string()
        progress["percentage"] = 20
        progress["step"] = "准备敏感性分析"

        # Step 1: 获取敏感性标准
        analysis_prompt = f"""
        请分析以下数据集的敏感性。字段：{', '.join(fields)}。
        样本数据：
        {sample_data}
        要求：
        1. 根据法律法规和安全性评估字段敏感性
        2. 为每个字段制定敏感性等级（高/中/低）并说明理由
        """
        standards = get_response(analysis_prompt, "sensitivity_standards")
        if not standards:
            raise ValueError("敏感性标准获取失败")
        parsed_standards = parse_response(standards, "sensitivity_standards")
        if not isinstance(parsed_standards, list):
            parsed_standards = [parsed_standards]
        progress["percentage"] = 40
        progress["step"] = "完成敏感性分析"

        # Step 2: 获取 AHP 矩阵
        matrix_prompt = f"""
        基于敏感性分析：{parsed_standards}
        请：
        1. 构建包含5个以上评价指标的AHP判断矩阵
        2. 指标涵盖数据安全、合规风险、业务影响等
        3. 使用1-9标度法
        4. 为每个指标提供解释
        """
        matrix_response = get_response(matrix_prompt, "ahp_matrix")
        if not matrix_response:
            raise ValueError("AHP矩阵获取失败")
        parsed_matrix = parse_response(matrix_response, "ahp_matrix")
        ahp_matrix = np.array(parsed_matrix["matrix"], dtype=float)
        indicators = {i["name"]: i["explanation"] for i in parsed_matrix["indicators"]}
        progress["percentage"] = 60
        progress["step"] = "完成 AHP 矩阵构建"

        # 计算权重
        weights = calculate_weights(ahp_matrix)
        is_consistent, cr_value = consistency_check(ahp_matrix)

        # Step 3: 获取字段评分
        scoring_prompt = f"""
        请为字段评分：
        - 字段: {fields}
        - 指标: {list(indicators.keys())}
        - 评分范围: 1-10
        要求：
        1. 每个指标单独评分
        2. 考虑敏感性和业务影响
        """
        scoring_response = get_response(scoring_prompt, "scoring")
        if not scoring_response:
            raise ValueError("评分数据获取失败")
        scoring_data = parse_response(scoring_response, "scoring")
        score_df = process_scoring_data(scoring_data, list(indicators.keys()))
        progress["percentage"] = 80
        progress["step"] = "完成字段评分"

        # 计算最终敏感度评分
        sensitivity_scores = score_df.values.dot(weights)

        # 生成 Excel 报告
        wb = Workbook()
        ws = wb.active
        ws.title = "敏感性分析报告"
        ws.append(["数据集", file_path])
        ws.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        ws.append([])
        ws.append(["敏感性初评"])
        ws.append(["数据", "敏感度", "理由"])
        for std in parsed_standards:
            ws.append([std["field"], std["sensitivity"], std["reason"]])
        ws.append([])
        ws.append(["数据评分"])
        header = ["数据"] + list(indicators.keys()) + ["综合评分"]
        ws.append(header)
        for idx, (field, row) in enumerate(score_df.iterrows()):
            score = sensitivity_scores[idx]
            ws.append([field] + list(row.values) + [score])
        red_fill = PatternFill(start_color="FF9999", end_color="FF9999", fill_type="solid")
        for row in ws[f"B5:B{ws.max_row}"]:
            for cell in row:
                if cell.value == "高":
                    cell.fill = red_fill
        wb.save(excel_report_path)
        logger.info(f"Excel 报告生成: {excel_report_path}")
        progress["step"] = "生成 Excel 报告"

        # 生成 Word 报告
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.add_heading("数据敏感性分析报告", 0).style.font.color.rgb = RGBColor(0, 0, 139)
        doc.add_heading("1. 敏感性标准", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "数据"
        hdr_cells[1].text = "敏感度"
        hdr_cells[2].text = "理由"
        for std in parsed_standards:
            row_cells = table.add_row().cells
            row_cells[0].text = std["field"]
            row_cells[1].text = std["sensitivity"]
            row_cells[2].text = std["reason"]
            if std["sensitivity"] == "高":
                row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
        doc.add_heading("2. 评价体系", level=1)
        doc.add_paragraph("AHP矩阵一致性：" + ("通过" if is_consistent else "未通过") + f" (CR={cr_value:.4f})")
        for ind, exp in indicators.items():
            doc.add_paragraph(f"{ind}: {exp}", style='List Bullet')
        doc.add_heading("3. 综合评分", level=1)
        table = doc.add_table(rows=1, cols=len(indicators) + 2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "字段"
        for i, ind in enumerate(indicators.keys(), 1):
            hdr_cells[i].text = ind
        hdr_cells[-1].text = "综合评分"
        for idx, (field, row) in enumerate(score_df.iterrows()):
            row_cells = table.add_row().cells
            row_cells[0].text = field
            for i, val in enumerate(row.values, 1):
                row_cells[i].text = f"{val:.2f}"
            row_cells[-1].text = f"{sensitivity_scores[idx]:.2f}"
        doc.save(word_report_path)
        logger.info(f"Word 报告生成: {word_report_path}")
        progress["percentage"] = 100
        progress["status"] = "done"
        progress["step"] = "完成"

    except Exception as e:
        logger.error(f"分析失败: {str(e)}")
        progress["status"] = "error"
        progress["step"] = f"错误: {str(e)}"


# 主页和上传
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return jsonify({"error": "No file uploaded"}), 400
        file = request.files['file']
        if not file.filename.endswith('.csv'):
            return jsonify({"error": "Please upload a CSV file"}), 400
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        excel_report_filename = f"report_{filename}_excel.xlsx"
        word_report_filename = f"report_{filename}_word.docx"
        excel_report_path = os.path.join(app.config['REPORT_FOLDER'], excel_report_filename)
        word_report_path = os.path.join(app.config['REPORT_FOLDER'], word_report_filename)

        global progress
        progress = {"status": "running", "percentage": 0, "step": "开始分析"}
        threading.Thread(target=analyze_and_generate_report,
                         args=(file_path, excel_report_path, word_report_path)).start()
        logger.info(f"文件上传成功，开始分析: {filename}")
        return jsonify({"message": "Analysis started"}), 200
    return render_template('index.html')


# 获取进度
@app.route('/progress')
def get_progress():
    return jsonify(progress)


# 展示结果
@app.route('/result')
def show_result():
    if progress["status"] != "done":
        return redirect(url_for('index'))

    # 查找最新的报告文件
    excel_files = [f for f in os.listdir(app.config['REPORT_FOLDER']) if f.endswith('_excel.xlsx')]
    word_files = [f for f in os.listdir(app.config['REPORT_FOLDER']) if f.endswith('_word.docx')]
    if not excel_files or not word_files:
        return redirect(url_for('index'))

    latest_excel = max(excel_files, key=lambda x: os.path.getctime(os.path.join(app.config['REPORT_FOLDER'], x)))
    latest_word = max(word_files, key=lambda x: os.path.getctime(os.path.join(app.config['REPORT_FOLDER'], x)))
    excel_path = os.path.join(app.config['REPORT_FOLDER'], latest_excel)
    word_path = os.path.join(app.config['REPORT_FOLDER'], latest_word)

    # 读取整个 Excel 文件以动态定位
    xl = pd.ExcelFile(excel_path)
    df = xl.parse(sheet_name=0, header=None)
    logger.info(f"Excel 文件完整内容:\n{df.to_string()}")  # 输出完整内容

    # 1. 数据集信息（前两行）
    dataset_info = df.iloc[:2, :2]  # 只取前两列，避免 NaN
    dataset_html = dataset_info.to_html(classes="table table-striped", index=False, header=False)

    # 2. 敏感性初评（修正为实际标题）
    sensitivity_start = None
    sensitivity_end = None
    if not df[df[0] == "敏感性初评"].empty:
        sensitivity_start = df[df[0] == "敏感性初评"].index[0] + 1
        sensitivity_end = df[df[0] == "数据评分"].index[0] - 1 if not df[df[0] == "数据评分"].empty else len(df)
        sensitivity_df = pd.read_excel(excel_path, skiprows=sensitivity_start,
                                       nrows=sensitivity_end - sensitivity_start, usecols=[0, 1, 2])
        sensitivity_html = sensitivity_df.to_html(classes="table table-striped", index=False)
    else:
        logger.warning("未找到 '敏感性初评'，尝试固定位置读取")
        sensitivity_df = pd.read_excel(excel_path, skiprows=4, nrows=21, usecols=[0, 1, 2])  # 假设21个字段
        if not sensitivity_df.empty and sensitivity_df.columns.tolist() == ["数据", "敏感度", "理由"]:
            sensitivity_html = sensitivity_df.to_html(classes="table table-striped", index=False)
        else:
            sensitivity_html = "<p>未找到敏感性初评数据</p>"

    # 3. 数据评分（修正为实际标题）
    scoring_start = None
    if not df[df[0] == "数据评分"].empty:
        scoring_start = df[df[0] == "数据评分"].index[0] + 1
        scoring_df = pd.read_excel(excel_path, skiprows=scoring_start)
        scoring_html = scoring_df.to_html(classes="table table-striped", index=False)
    else:
        logger.warning("未找到 '数据评分'，尝试固定位置读取")
        scoring_df = pd.read_excel(excel_path, skiprows=27)  # 假设从第28行开始
        if not scoring_df.empty and "数据" in scoring_df.columns:
            scoring_html = scoring_df.to_html(classes="table table-striped", index=False)
        else:
            scoring_html = "<p>未找到数据评分数据</p>"

    # 处理 Word 文件
    with open(word_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        word_html = result.value

    return render_template('result.html', dataset_html=dataset_html, sensitivity_html=sensitivity_html,
                           scoring_html=scoring_html, word_html=word_html,
                           excel_filename=latest_excel, word_filename=latest_word)


# 下载文件
@app.route('/download/<filename>')
def download_report(filename):
    return send_from_directory(app.config['REPORT_FOLDER'], filename)



# 本地运行调试
if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 10000))  # Render 会自动注入 PORT 环境变量
    app.run(host='0.0.0.0', port=port)
